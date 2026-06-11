"""Build PEAP-AI technical summary as a Word document (.docx)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_DOCX = DOCS / "PEAP-AI-Technical-Summary.docx"
GRAPH_PNG = DOCS / "peap-ai-feedback-graph.png"
GRAPH_MMD = DOCS / "peap-ai-feedback-graph.mmd"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value


def build_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("PEAP-AI Technical Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This document summarizes the workflow, architecture, programming stack, and "
        "LLM usage of PEAP-AI as implemented in the LLED_bot_MVP repository "
        "(LLED 200 Essay Feedback MVP) for formative AI feedback on Descriptive Report drafts."
    )

    add_heading(doc, "1. Workflow and Architecture", 1)

    add_heading(doc, "1.1 System Overview", 2)
    doc.add_paragraph(
        "PEAP-AI is a locally hosted web application MVP. A student submits an essay draft; "
        "the backend orchestrates multiple LLM calls through a LangGraph StateGraph; the "
        "system returns strictly schema-validated JSON feedback; the browser renders a "
        "two-column view with paragraph-level annotations, sidebar cards, and an overall summary."
    )
    doc.add_paragraph(
        "Current scope excludes user authentication, persistence, and RAG vector retrieval. "
        "Course material references are attached via deterministic mapping, not retrieval."
    )

    add_heading(doc, "1.2 End-to-End Workflow", 2)
    add_table(
        doc,
        ["Stage", "Behavior"],
        [
            [
                "Input",
                "Paste plain text or upload .txt, .md, .doc, .docx, or .pdf; "
                "fileExtractor extracts essay_text.",
            ],
            [
                "Submit",
                "POST /api/essay-feedback with body { essay_text }.",
            ],
            [
                "Orchestration",
                "runFeedbackGraph: parse paragraphs, load prompt.md, run three parallel "
                "dimension branches, merge (up to 12 annotations), validate with Zod, "
                "optionally repair JSON once, attach course material labels.",
            ],
            [
                "Output",
                "FeedbackResponse with paragraphs, annotations (offsets, function, level, "
                "evidence, revision guidance), and overall_feedback.",
            ],
            [
                "UI",
                "Highlights linked to sidebar cards, dimension/level filters, Summary tab, "
                "sessionStorage cache for the latest result.",
            ],
        ],
    )

    add_heading(doc, "1.3 LangGraph Feedback Workflow (Implemented)", 2)
    doc.add_paragraph(
        "The diagram below reflects the StateGraph compiled in "
        "src/backend/lib/feedbackGraph.ts (@langchain/langgraph)."
    )
    if GRAPH_PNG.exists():
        doc.add_picture(str(GRAPH_PNG), width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(
            "[Diagram image not found. Regenerate with: "
            "npx @mermaid-js/mermaid-cli -i docs/peap-ai-feedback-graph.mmd "
            "-o docs/peap-ai-feedback-graph.png]"
        )

    doc.add_paragraph("Graph nodes (in execution order):")
    nodes = [
        "prepare_context — parseEssay + buildPromptMessages from prompt.md",
        "content_feedback — parallel LLM branch (Content function, max 4 annotations)",
        "interpersonal_feedback — parallel LLM branch (Interpersonal function)",
        "organization_feedback — parallel LLM branch (Organization function)",
        "merge_feedback — merge branches, sort by severity, cap at 12 annotations",
        "validate_feedback — Zod validation of offsets, paragraph IDs, and schema",
        "repair_feedback (conditional) — one repair attempt if validation fails",
        "attach_course_materials — map function + level to course material labels",
        "return_feedback — return FeedbackResponse to the API layer",
    ]
    for node in nodes:
        add_bullet(doc, node)

    doc.add_paragraph(
        "Typical LLM calls per submission: about 3 (parallel dimension nodes). "
        "If repair runs, add 1 additional call."
    )

    add_heading(doc, "1.4 Layered Architecture", 2)
    layers = [
        ("Presentation", "src/frontend/ — HTML, CSS, TypeScript"),
        ("API", "src/backend/server.ts — Express 5"),
        ("Orchestration", "src/backend/lib/feedbackGraph.ts — LangGraph StateGraph"),
        (
            "LLM + Prompt",
            "openaiClient.ts, promptBuilder.ts, prompts/prompt.md",
        ),
        ("Validation + Shared", "schemaValidator.ts (Zod), src/shared/schema.ts"),
        ("Course mapping", "src/shared/courseMaterials.ts"),
    ]
    add_table(doc, ["Layer", "Location / role"], layers)

    add_heading(doc, "1.5 Evaluation Framework", 2)
    doc.add_paragraph(
        "Feedback aligns with the Academic Writing Matrix: three functional dimensions "
        "(Content, Interpersonal, Organization) crossed with three levels "
        "(text, section, clause_word). The prompt forbids scoring and sentence rewriting; "
        "it provides diagnostic feedback and revision guidance only."
    )

    add_heading(doc, "2. Programming Languages and Stack", 1)
    add_table(
        doc,
        ["Category", "Choice"],
        [
            ["Primary language", "TypeScript (strict: true) across backend, frontend, shared"],
            ["Markup & styling", "HTML5 and CSS (no React/Vue SPA framework)"],
            ["Runtime", "Node.js 20 or newer"],
            ["Backend", "Express 5"],
            ["Workflow engine", "LangGraph (@langchain/langgraph)"],
            ["Validation", "Zod 4"],
            ["Document parsing", "mammoth, pdf-parse, word-extractor"],
            ["Uploads", "multer (in-memory, max 10 MB per file)"],
            ["Build", "tsc for frontend bundle; tsx watch + concurrently in dev"],
        ],
    )

    add_heading(doc, "3. LLM Models", 1)
    doc.add_paragraph(
        "PEAP-AI currently uses OpenAI models as the primary integration path. The backend "
        "adapter (openaiClient.ts) calls the OpenAI Chat Completions API with structured "
        "JSON schema output (response_format: json_schema, strict: true)."
    )
    doc.add_paragraph(
        "The architecture is adapter-based: additional large language models (e.g., Anthropic, "
        "Google, Azure OpenAI, or other OpenAI-compatible endpoints) can be wired in by "
        "implementing the same structured-feedback contract behind the LangGraph nodes, "
        "without changing the frontend schema."
    )
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Primary provider", "OpenAI (current production path)"],
            [
                "Default model",
                "gpt-5.4-mini (overridable via OPENAI_MODEL; see .env.example)",
            ],
            ["Configuration", "OPENAI_API_KEY (required), OPENAI_MODEL (optional), PORT"],
            [
                "Invocation pattern",
                "Parallel dimension nodes + optional single repair pass",
            ],
            [
                "Extensibility",
                "Other LLMs supported in principle via adapter swap or multi-provider routing",
            ],
        ],
    )

    add_heading(doc, "4. Document Metadata", 1)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Title", "PEAP-AI Technical Summary"],
            ["Version", "MVP (essay-feedback-mvp v0.1.0)"],
            ["Date", date.today().isoformat()],
        ],
    )

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("Prepared by: ").bold = True
    sig.add_run("Lin Manshan (林满山)\n")
    sig.add_run("Email: lin.manshan050813@gmail.com")

    doc.add_paragraph(
        "Source: LLED_bot_MVP repository — for course documentation and future HelpMe migration."
    )

    return doc


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
